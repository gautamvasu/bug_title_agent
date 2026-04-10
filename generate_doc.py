"""Generate DefectLens Architecture Document in DOCX format."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 51, 102)

# -- Title Page --
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DefectLens')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Architecture & User Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Internal Tool\n').font.size = Pt(12)
meta.add_run('Powered by MetaGen & Llama\n\n').font.size = Pt(12)
meta.add_run('Version: 1.0\n').font.size = Pt(11)
meta.add_run('April 2026').font.size = Pt(11)

doc.add_page_break()

# -- Table of Contents --
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Architecture',
    '   2.1 System Architecture',
    '   2.2 Component Details',
    '   2.3 Data Flow',
    '   2.4 AI Provider Integration',
    '3. Installation',
    '   3.1 Prerequisites',
    '   3.2 Setup Steps',
    '   3.3 AI Provider Configuration',
    '4. Usage Guide',
    '   4.1 Review Modes',
    '   4.2 Checklist & Mandatory Tags',
    '   4.3 Log Parsing',
    '   4.4 Review Output',
    '   4.5 Google Chat Notifications',
    '5. Technical Reference',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# -- 1. Overview --
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'DefectLens is a bug report review tool that helps QA engineers and developers '
    'evaluate the completeness of Phabricator tasks. It uses AI (MetaGen/Llama) to analyze bug reports '
    'against configurable checklists, verify mandatory tags, parse log files for error signals, '
    'and suggest improved defect titles.'
)
doc.add_paragraph('Key capabilities:', style='List Bullet')
features = [
    'Checklist Gap Analysis — Compare bug reports against a configurable checklist with color-coded results (green=present, amber=partial, red=missing).',
    'Mandatory Tag Verification — Check required tags with wildcard support (e.g., FoundBy*, *testing).',
    'Log Parsing — Extract crashes, ANRs, exceptions, and tombstones from bugreport/logcat files.',
    'Defect Title Suggestions — AI-generated title improvements based on description and log signals.',
    'Multi-Task Review — Review single tasks, multiple tasks, or all open tasks by creator with date filtering.',
    'Closed Task Detection — Automatically blocks review of closed/resolved tasks.',
    'Google Chat Notifications — Send review results to task creators via Google Chat DM.',
    'Auto-fetch from Phabricator — Pulls task title, description, creator, status, and tags via jf CLI.',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet 2')

# -- 2. Architecture --
doc.add_heading('2. Architecture', level=1)

doc.add_heading('2.1 System Architecture', level=2)
doc.add_paragraph(
    'DefectLens is a single-page Streamlit web application that integrates with multiple internal '
    'services and AI providers. The application runs locally on the user\'s machine or on a shared devserver.'
)

# Architecture diagram as text
doc.add_paragraph()
arch = doc.add_paragraph()
arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = arch.add_run(
    '┌─────────────────────────────────────────────────┐\n'
    '│                  DefectLens UI                   │\n'
    '│              (Streamlit Web App)                 │\n'
    '├─────────────────────────────────────────────────┤\n'
    '│                                                 │\n'
    '│  ┌──────────┐  ┌──────────┐  ┌──────────────┐  │\n'
    '│  │ Task     │  │ Review   │  │ Notification │  │\n'
    '│  │ Fetcher  │  │ Engine   │  │ Module       │  │\n'
    '│  └────┬─────┘  └────┬─────┘  └──────┬───────┘  │\n'
    '│       │             │               │           │\n'
    '├───────┼─────────────┼───────────────┼───────────┤\n'
    '│       │             │               │           │\n'
    '│       ▼             ▼               ▼           │\n'
    '│  Phabricator   MetaGen/Ollama   Google Chat     │\n'
    '│  (jf graphql)  (Llama API)      (gchat CLI)     │\n'
    '└─────────────────────────────────────────────────┘\n'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('2.2 Component Details', level=2)

components = [
    ('Streamlit UI (app.py)',
     'The main application file containing all UI logic, API integrations, and data processing. '
     'Built with Streamlit framework for rapid prototyping of data-centric web apps. '
     'Handles three review modes: Single Task, Multiple Tasks, and All Open Tasks by Creator.'),
    ('Task Fetcher',
     'Interfaces with Phabricator via the jf graphql CLI. Fetches task details including '
     'title, description, creator, tags, and status. For batch mode, uses PowerSearch API '
     '(task_search_query) to find open tasks by creator with date range filtering.'),
    ('Review Engine',
     'Core analysis module that orchestrates the review process:\n'
     '• Mandatory Tag Check — Compares task tags against required tags with wildcard matching.\n'
     '• Log Parser — Regex-based extraction of FATAL exceptions, ANRs, native crashes, tombstones, and exceptions from bugreport/logcat files.\n'
     '• AI Analysis — Sends task context (title, description, logs, checklist, tags) to the LLM for gap analysis and title suggestions.\n'
     '• Result Colorizer — Converts LLM output to styled HTML with color-coded status indicators.'),
    ('Notification Module',
     'Sends review results to task creators via Google Chat using the gchat CLI at '
     '/opt/facebook/bin/gchat. Messages are written to a temp file and sent via --text-file flag '
     'to handle long content.'),
    ('AI Provider Layer',
     'Supports two AI providers:\n'
     '• MetaGen (Internal) — Uses the Llama API at api.llama.com/compat/v1/chat/completions with Llama-4-Scout model. Requires a MetaGen API key.\n'
     '• Ollama (Local) — Runs Llama 3.1 8B locally via Ollama. No API key needed, fully offline.'),
]

for name, desc in components:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

doc.add_heading('2.3 Data Flow', level=2)
steps = [
    'User enters task number(s) or creator unixname.',
    'Task Fetcher queries Phabricator GraphQL API for task details (title, description, creator, tags, status).',
    'User optionally uploads log files, provides checklist, and defines mandatory tags.',
    'On "Review Task" click, mandatory tags are checked locally (no AI needed).',
    'Task context is assembled into a structured prompt and sent to the AI provider (MetaGen or Ollama).',
    'AI returns a structured review with checklist gap analysis and title suggestions.',
    'Result is parsed and colorized into styled HTML for display.',
    'User can optionally notify the task creator on Google Chat.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('2.4 AI Provider Integration', level=2)

# MetaGen table
doc.add_paragraph('MetaGen (Internal):', style='List Bullet')
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
cells = [
    ('Endpoint', 'https://api.llama.com/compat/v1/chat/completions'),
    ('Model', 'Llama-4-Scout-17B-16E-Instruct-FP8'),
    ('Auth', 'Bearer token (MetaGen API key)'),
    ('Protocol', 'OpenAI-compatible Chat Completions API'),
    ('Key Portal', 'https://metagen-llm-api-keys.nest.x2p.facebook.net/'),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph('Ollama (Local):', style='List Bullet')
table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Light Grid Accent 1'
cells2 = [
    ('Endpoint', 'http://localhost:11434/api/chat'),
    ('Model', 'llama3.1:8b'),
    ('Auth', 'None (local)'),
    ('Protocol', 'Ollama Chat API'),
]
for i, (k, v) in enumerate(cells2):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v

doc.add_page_break()

# -- 3. Installation --
doc.add_heading('3. Installation', level=1)

doc.add_heading('3.1 Prerequisites', level=2)
prereqs = [
    'Python 3.9 or higher',
    'jf CLI (for Phabricator integration)',
    'gchat CLI at /opt/facebook/bin/gchat (for Google Chat notifications)',
    'VPN (for MetaGen API access) or Ollama installed locally',
]
for p in prereqs:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('3.2 Setup Steps', level=2)

doc.add_paragraph('Step 1: Clone the repository')
code = doc.add_paragraph()
run = code.add_run('git clone https://github.com/gautamvasu/defectlens.git\ncd defectlens')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph('Step 2: Create a virtual environment and install dependencies')
code = doc.add_paragraph()
run = code.add_run('python3 -m venv venv\nsource venv/bin/activate\npip install -r requirements.txt')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph('Step 3: Start the application')
code = doc.add_paragraph()
run = code.add_run('streamlit run app.py --server.headless true')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph('The app opens at http://localhost:8501')

doc.add_heading('3.3 AI Provider Configuration', level=2)

doc.add_paragraph('Option A: MetaGen (Internal — recommended)')
metagen_steps = [
    'Connect to VPN.',
    'Visit https://metagen-llm-api-keys.nest.x2p.facebook.net/ and click "Create API Key".',
    'Go to the MetaGen entitlements portal and create a new entitlement (e.g., "your_unixname_key").',
    'Link your Llama API App ID to the MetaGen entitlement.',
    'Either paste the key in the sidebar when running the app, or add it to the .env file:',
]
for s in metagen_steps:
    doc.add_paragraph(s, style='List Number')
code = doc.add_paragraph()
run = code.add_run('METAGEN_API_KEY=your-key-here')
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Option B: Ollama (Local — no API key needed)')
ollama_steps = [
    'Download and install Ollama from ollama.com/download.',
    'Open the Ollama app.',
    'Pull the model by running in terminal: ollama pull llama3.1:8b',
    'Make sure Ollama is running before using DefectLens.',
    'Select "Ollama (Local — No API Key)" from the sidebar dropdown.',
]
for s in ollama_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# -- 4. Usage Guide --
doc.add_heading('4. Usage Guide', level=1)

doc.add_heading('4.1 Review Modes', level=2)

doc.add_heading('Single Task', level=3)
doc.add_paragraph(
    'Enter a single task number (e.g., T12345). DefectLens automatically fetches the task title, '
    'description, creator, tags, and status from Phabricator. If the task is closed, a warning is shown '
    'and the review is blocked. You can manually edit the title and description before reviewing.'
)

doc.add_heading('Multiple Tasks', level=3)
doc.add_paragraph(
    'Enter multiple task numbers, one per line. Each task is fetched and reviewed independently. '
    'A progress bar shows the review status. Closed or not-found tasks are skipped with a reason. '
    'Results are displayed in expandable sections, one per task.'
)

doc.add_heading('All Open Tasks by Creator', level=3)
doc.add_paragraph(
    'Enter a creator\'s unixname to fetch all their open tasks from Phabricator. '
    'Filter by date range (7, 14, 30, 60, 90, 180 days, or all time) and set a maximum task count (5-50). '
    'Tasks appear as checkboxes — use "Select all" to toggle all, or pick individual tasks. '
    'Only selected tasks are reviewed when you click the review button.'
)

doc.add_heading('4.2 Checklist & Mandatory Tags', level=2)

doc.add_paragraph('Checklist (optional):', style='List Bullet')
doc.add_paragraph(
    'Provide a list of required information items for a complete bug report. '
    'The AI compares each checklist item against the task and reports whether it is '
    'PRESENT (green), PARTIALLY PRESENT (amber), or MISSING (red). '
    'Input methods: paste manually, upload a file (.txt, .csv, .xlsx), or link a Google Sheet.'
)

doc.add_paragraph('Mandatory Tags (optional):', style='List Bullet')
doc.add_paragraph(
    'Define tags that must be present on every task. DefectLens checks the task\'s actual tags '
    'against your list and reports matches/misses. Supports wildcard matching:\n'
    '• FoundBy* — matches any tag starting with "FoundBy" (e.g., FoundBy-QA, FoundBy-User)\n'
    '• *testing — matches any tag ending with "testing"\n'
    'Input methods: paste manually, upload a file, or link a Google Sheet.'
)

doc.add_heading('4.3 Log Parsing', level=2)
doc.add_paragraph(
    'Upload a bugreport or logcat file (.txt, .log, .gz, .zip). DefectLens parses it using regex '
    'patterns to extract up to 10 unique error signals:'
)
log_types = [
    'FATAL — Fatal exceptions, AndroidRuntime crashes, kernel panics',
    'ANR — Application Not Responding events, input dispatching timeouts',
    'NATIVE_CRASH — Native crash dumps, tombstone markers',
    'TOMBSTONE — Signal information, abort messages, backtraces',
    'EXCEPTION — Java/Kotlin exceptions with stack traces',
]
for lt in log_types:
    doc.add_paragraph(lt, style='List Bullet')
doc.add_paragraph(
    'Extracted signals are included in the AI prompt for more accurate title suggestions '
    'and gap analysis.'
)

doc.add_heading('4.4 Review Output', level=2)
doc.add_paragraph('The review output is divided into sections with color-coded results:')
output_items = [
    'Mandatory Tags — Score line (color-coded by percentage: green >= 80%, amber 50-79%, red < 50%) followed by individual tag results.',
    'Checklist Gap Analysis — Overall completeness score followed by per-item status with color coding.',
    'Suggested Defect Title — Analysis of the current title, three ranked suggestions in blue, and explanation of the top pick.',
]
for item in output_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.5 Google Chat Notifications', level=2)
doc.add_paragraph(
    'After reviewing a task, click the green "Notify [Creator] on Google Chat" button to send the '
    'review results as a plain-text DM to the task creator. The message includes the full review '
    'with all gaps and suggestions. For multi-task reviews, each task has its own notify button.'
)

doc.add_page_break()

# -- 5. Technical Reference --
doc.add_heading('5. Technical Reference', level=1)

doc.add_paragraph('File Structure:', style='List Bullet')
files = doc.add_table(rows=5, cols=2)
files.style = 'Light Grid Accent 1'
file_data = [
    ('app.py', 'Main Streamlit application — UI, API integrations, review logic'),
    ('defectlens_cli.py', 'CLI version for command-line usage'),
    ('requirements.txt', 'Python dependencies (streamlit, pandas, openpyxl, python-dotenv, markdown)'),
    ('.env', 'Environment variables (MetaGen API key)'),
    ('.gitignore', 'Git ignore rules (.env, venv/, __pycache__/, .DS_Store)'),
]
for i, (k, v) in enumerate(file_data):
    files.rows[i].cells[0].text = k
    files.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph('Key Functions:', style='List Bullet')
funcs = doc.add_table(rows=10, cols=2)
funcs.style = 'Light Grid Accent 1'
func_data = [
    ('fetch_task_details()', 'Fetch task name, description, creator, tags, status from Phabricator'),
    ('fetch_user_fbid()', 'Resolve unixname to FBID for task search'),
    ('fetch_open_tasks_by_owner()', 'Search open tasks by creator with date range filter'),
    ('check_mandatory_tags()', 'Compare mandatory tags (with wildcards) against task tags'),
    ('parse_log()', 'Extract error signals from bugreport/logcat files'),
    ('build_user_prompt()', 'Assemble the LLM prompt with all task context'),
    ('call_metagen()', 'Send review request to MetaGen Llama API'),
    ('call_ollama()', 'Send review request to local Ollama instance'),
    ('colorize_result()', 'Convert LLM output to color-coded HTML'),
    ('send_gchat_message()', 'Send DM via gchat CLI'),
]
for i, (k, v) in enumerate(func_data):
    funcs.rows[i].cells[0].text = k
    funcs.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph('External Dependencies:', style='List Bullet')
deps = doc.add_table(rows=4, cols=2)
deps.style = 'Light Grid Accent 1'
dep_data = [
    ('jf graphql', 'Phabricator GraphQL API access — task details and PowerSearch'),
    ('gchat CLI', 'Google Chat messaging at /opt/facebook/bin/gchat'),
    ('MetaGen API', 'Llama LLM inference at api.llama.com'),
    ('Ollama', 'Local LLM inference at localhost:11434'),
]
for i, (k, v) in enumerate(dep_data):
    deps.rows[i].cells[0].text = k
    deps.rows[i].cells[1].text = v

# -- Save --
output_path = '/Users/vgautam/defectlens/DefectLens_Architecture_Document.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
