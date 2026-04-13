#!/usr/bin/env python3
"""Convert PRD.md to PRD.docx with proper formatting."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# -- Title --
title = doc.add_heading('DefectLens — Product Requirements Document', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -- Metadata --
meta = [
    ('Author', 'Vasu Gautam'),
    ('Date', 'April 2026'),
    ('Status', 'Active'),
    ('Version', '1.0'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(value)

doc.add_paragraph()  # spacer

# -- Helper to add a table --
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()  # spacer after table


# ========== Section 1: Overview ==========
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'DefectLens is an AI-powered bug report review tool that helps QA engineers and developers '
    'write higher-quality defect reports. It analyzes Phabricator tasks against configurable '
    'checklists, verifies mandatory tags, parses log files for error signals, and suggests '
    'clearer defect titles — all through a web-based interface powered by Streamlit.'
)

# ========== Section 2: Problem Statement ==========
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph('Bug reports filed in Phabricator often suffer from:')
problems = [
    ('Incomplete information', 'missing reproduction steps, expected vs. actual behavior, environment details, or log attachments'),
    ('Poor titles', 'vague titles like "login not working" that don\'t convey the component, symptom, or context'),
    ('Missing tags', 'mandatory tags (severity, platform, found-by source) are frequently omitted, making triage and tracking harder'),
    ('No feedback loop', 'creators don\'t learn what\'s missing until a reviewer manually points it out, often days later'),
]
for title_text, desc in problems:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title_text} — ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'These gaps slow down triage, increase back-and-forth, and reduce the signal-to-noise ratio in defect backlogs.'
)

# ========== Section 3: Goals ==========
doc.add_heading('3. Goals', level=1)
add_table(
    ['Goal', 'Success Metric'],
    [
        ['Improve bug report completeness', 'Increase average checklist coverage score across reviewed tasks'],
        ['Reduce review turnaround time', 'Provide instant automated review instead of waiting for manual reviewer'],
        ['Standardize defect quality', 'Consistent checklist and tag enforcement across teams'],
        ['Close the feedback loop', 'Notify creators immediately via Google Chat with actionable gaps'],
    ]
)

# ========== Section 4: Non-Goals ==========
doc.add_heading('4. Non-Goals', level=1)
for item in [
    'Automatically editing or closing Phabricator tasks',
    'Replacing human code review or manual QA judgment',
    'Providing root-cause analysis or debugging suggestions',
    'Supporting non-Phabricator issue trackers',
]:
    doc.add_paragraph(item, style='List Bullet')

# ========== Section 5: Target Users ==========
doc.add_heading('5. Target Users', level=1)
add_table(
    ['Persona', 'Use Case'],
    [
        ['QA Engineers', 'Review their own or teammates\' bug reports before handoff to engineering'],
        ['QA Leads / Managers', 'Batch-review all open tasks by a specific creator to audit report quality'],
        ['Developers', 'Quickly assess incoming bug reports for completeness before investing in investigation'],
    ]
)

# ========== Section 6: Features & Requirements ==========
doc.add_heading('6. Features & Requirements', level=1)

# 6.1
doc.add_heading('6.1 Review Modes', level=2)
add_table(
    ['Mode', 'Description', 'Priority'],
    [
        ['Single Task', 'Enter one task number; auto-fetches title, description, creator, tags, and status from Phabricator', 'P0'],
        ['Multiple Tasks', 'Paste multiple task numbers (one per line); each is reviewed independently with a progress bar', 'P0'],
        ['All Open Tasks by Creator', 'Enter a unixname, filter by date range (7-180 days or all time), select tasks, and batch-review', 'P1'],
    ]
)

# 6.2
doc.add_heading('6.2 Checklist Gap Analysis', level=2)
for item in [
    'Compare the bug report\'s content against a configurable checklist',
    'Each item is rated: PRESENT (green), PARTIALLY PRESENT (amber), or MISSING (red)',
    'An overall completeness score (X/Y items, Z%) is displayed',
    'Checklist input methods: paste manually, upload file (.txt, .csv, .xlsx), or link a Google Sheet',
    'If no checklist is provided, a default assessment covers: steps to reproduce, expected vs. actual behavior, environment info, severity, log attachment, and screenshots',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6.3
doc.add_heading('6.3 Mandatory Tag Verification', level=2)
for item in [
    'Define required tags that must be applied to every task',
    'Supports wildcard matching (FoundBy* matches any tag starting with "FoundBy"; *testing matches any tag ending with "testing")',
    'Each tag is reported as PRESENT or MISSING',
    'A mandatory tags score (X/Y present, Z%) is displayed',
    'Tag input methods: paste manually, upload file, or link a Google Sheet',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6.4
doc.add_heading('6.4 Log Parsing', level=2)
for item in [
    'Upload bugreport or logcat files (.txt, .log, .gz, .zip)',
    'Automatically extracts and categorizes: FATAL exceptions, ANRs, native crashes, tombstones, and general exceptions',
    'Provides context (surrounding lines) for each signal',
    'Deduplicates signals and caps at 10 to keep LLM context manageable',
    'Parsed signals are included in the AI review prompt for better title suggestions',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6.5
doc.add_heading('6.5 Defect Title Suggestions', level=2)
for item in [
    'Generates 3 ranked title suggestions based on the task description, current title, and parsed log signals',
    'Each suggestion follows best practices: clear, concise (<80 chars), specific component + symptom + context',
    'Includes analysis of the current title\'s weaknesses and rationale for the top suggestion',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6.6
doc.add_heading('6.6 Closed Task Detection', level=2)
for item in [
    'Automatically checks task status via Phabricator API',
    'Blocks review of closed tasks with a warning message',
    'In batch mode, closed tasks are skipped and listed separately',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6.7
doc.add_heading('6.7 Google Chat Notification', level=2)
for item in [
    'One-click notification to send review results to the task creator via Google Chat DM',
    'Message is formatted for gchat (bold, plain text, no HTML)',
    'Uses the gchat CLI at /opt/facebook/bin/gchat',
    'Available per-task in both single and multi-task review modes',
]:
    doc.add_paragraph(item, style='List Bullet')

# 6.8
doc.add_heading('6.8 CLI Tool', level=2)
for item in [
    'Standalone CLI (defectlens_cli.py) for quick title suggestions without the web UI',
    'Supports single-shot (defectlens_cli.py <task> <title>) and interactive modes',
    'Uses Claude API (Anthropic SDK) for generation',
]:
    doc.add_paragraph(item, style='List Bullet')

# ========== Section 7: Architecture ==========
doc.add_heading('7. Architecture', level=1)

doc.add_heading('7.1 Components', level=2)
doc.add_paragraph(
    'The system consists of a Streamlit web UI that communicates with the core logic in app.py. '
    'app.py integrates with Phabricator (via jf CLI) for task data, MetaGen or Ollama for AI inference, '
    'and the gchat CLI for notifications.'
)

doc.add_heading('7.2 AI Providers', level=2)
add_table(
    ['Provider', 'Model', 'Use Case'],
    [
        ['MetaGen (Internal)', 'Llama-4-Scout-17B-16E-Instruct-FP8', 'Primary — requires VPN and API key'],
        ['Ollama (Local)', 'llama3.1:8b', 'Offline / no API key — runs on local machine'],
    ]
)

doc.add_heading('7.3 External Dependencies', level=2)
add_table(
    ['Dependency', 'Purpose'],
    [
        ['jf CLI', 'Phabricator GraphQL queries (task details, user lookup, task search)'],
        ['gchat CLI', 'Google Chat DM notifications'],
        ['MetaGen API', 'LLM inference (internal)'],
        ['Ollama', 'LLM inference (local)'],
    ]
)

# ========== Section 8: Tech Stack ==========
doc.add_heading('8. Tech Stack', level=1)
add_table(
    ['Layer', 'Technology'],
    [
        ['Frontend', 'Streamlit 1.56+'],
        ['Language', 'Python 3.9+'],
        ['AI/LLM', 'MetaGen (Llama) / Ollama'],
        ['Data', 'Phabricator (via jf GraphQL)'],
        ['Notifications', 'Google Chat (via gchat CLI)'],
        ['Dependencies', 'streamlit, pandas, openpyxl, python-dotenv, markdown'],
    ]
)

# ========== Section 9: User Flow ==========
doc.add_heading('9. User Flow', level=1)
steps = [
    'Select review mode — Single Task, Multiple Tasks, or All Open Tasks by Creator',
    'Enter task number(s) — Task details are auto-fetched from Phabricator',
    'Upload log file (optional) — Bugreport/logcat parsed for error signals',
    'Provide checklist (optional) — Paste, upload, or link a Google Sheet',
    'Define mandatory tags (optional) — Paste, upload, or link a Google Sheet',
    'Click "Review Task(s)" — AI analyzes the task and returns mandatory tag check, checklist gap analysis, and 3 suggested defect titles',
    'Notify creator (optional) — Send results to the task creator on Google Chat',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{step}', style='List Number')

# ========== Section 10: Constraints & Assumptions ==========
doc.add_heading('10. Constraints & Assumptions', level=1)
for item in [
    'Users have access to Meta\'s internal network (VPN) for MetaGen and Phabricator',
    'The jf CLI is installed and authenticated',
    'The gchat CLI is available at /opt/facebook/bin/gchat for notifications',
    'LLM responses are capped at 1024 tokens to keep reviews focused',
    'Log parsing extracts at most 10 signals to stay within LLM context limits',
    'Google Sheets used for checklists/tags must be shared/accessible',
]:
    doc.add_paragraph(item, style='List Bullet')

# ========== Section 11: Future Considerations ==========
doc.add_heading('11. Future Considerations', level=1)
futures = [
    ('Auto-attach review to Phabricator task', 'Post the review as a comment on the task itself'),
    ('Team dashboards', 'Aggregate completeness scores across teams over time'),
    ('Custom scoring weights', 'Allow teams to weight checklist items differently'),
    ('Integration with CI/CD', 'Trigger reviews automatically when tasks are created or updated'),
    ('Historical trend analysis', 'Track improvement in bug report quality per creator over time'),
    ('Additional notification channels', 'Workplace, email, or Phabricator Herald rules'),
]
for title_text, desc in futures:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title_text} — ')
    run.bold = True
    p.add_run(desc)

# -- Save --
output_path = '/Users/vgautam/defectlens/DefectLens_PRD.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
